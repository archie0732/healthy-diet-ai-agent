import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
EXP = ROOT / "experiments" / "version_aware_rag"
PACKET = (
    EXP
    / "data"
    / "annotations_v5"
    / "r2_22_gpt56_blind_review"
    / "BLIND_PACKET.jsonl"
)
ZH_TRANSLATIONS = (
    EXP
    / "data"
    / "annotations_v5"
    / "r2_22_gpt56_blind_review"
    / "PASSAGE_TRANSLATIONS_ZH.jsonl"
)
OUT = EXP / "reviewer_packets"

ZH_QUESTIONS = {
    "blind-01": "增加鉀攝取量對不同基線攝取量與是否患有高血壓的族群，其血壓影響為何？",
    "blind-02": "在不同鉀與鈉暴露量下，證據如何描述鉀對血壓的影響？",
    "blind-03": "使用非糖甜味劑控制體重時，試驗結果顯示什麼？哪些比較情境會限制結果的解讀？",
    "blind-04": "應如何解讀孕期使用非糖甜味劑的研究結果，包括觀察到的結果與證據確定性？",
    "blind-05": "制定政策時，應如何同時考量減鈉的益處與低鈉鹽替代品的風險？",
    "blind-06": "人口層級的減鹽計畫應依據哪些鈉攝取目標證據與政策條件？",
    "blind-07": "身體活動應如何廣泛定義？多元組合的活動計畫可以包含哪些內容？",
    "blind-08": "選擇有助於睡眠與慢性病效益的活動時，應考量哪些證據與不確定性？",
    "blind-09": "制定總脂肪攝取建議時，應如何考量證據強度與營養素替代效應？",
    "blind-10": "執行總脂肪政策時，需要哪些監測措施與營養保障？",
    "blind-11": "對多數人而言，適當的鉀來源有哪些？哪些族群需要特別注意鉀暴露？",
    "blind-12": "飲用水中的鉀暴露與食物來源的可行性，應如何納入人口鉀攝取規劃？",
    "blind-13": "介入研究與世代研究顯示，非糖甜味劑對血糖、非傳染性疾病與死亡結果有何影響？",
    "blind-14": "應如何解讀非糖甜味劑試驗研究與世代研究結果之間的差異？",
    "blind-15": "哪些認知效益支持對身心障礙者提出身體活動指引？",
    "blind-16": "評估成人久坐行為時，會考量哪些結果與效果修飾因子？",
    "blind-17": "使用低鈉鹽替代品時，應同時提供哪些安全性證據？",
    "blind-18": "針對家庭的減鹽介入措施，有哪些證據限制與可接受性考量？",
    "blind-19": "公共衛生目標與健康公平考量應如何影響總脂肪政策？",
    "blind-20": "人口營養目標與總脂肪指引的適用範圍有何關係？",
    "blind-21": "儘管正常血壓族群仍有證據缺口，鉀指引為何仍認為廣泛增加鉀攝取具有益處？",
    "blind-22": "哪些甜味劑屬於非糖甜味劑建議的適用範圍？哪些不屬於？",
    "blind-23": "孕期與產後身體活動建議涵蓋哪些人，以及哪些活動領域？",
    "blind-24": "低鈉鹽替代品的證據對心血管事件顯示什麼結果？",
    "blind-25": "WHO 對不同年齡層建議或提議的天然膳食纖維攝取量為何？",
    "blind-26": "針對總脂肪與預防不健康體重增加，證據回顧的範圍如何界定？",
    "blind-27": "計畫規劃者應如何在公共衛生營養政策中使用 WHO 鉀攝取指引？",
    "blind-28": "指引如何描述非糖甜味劑世代研究中，為處理反向因果關係所進行的敏感度分析？",
    "blind-29": "在什麼情況下，運動前通常不需要事先取得醫療許可？",
    "blind-30": "WHO 低鈉鹽替代品指引所陳述的目標與預期政策用途為何？",
    "blind-31": "WHO 碳水化合物指引對兩歲以下兒童的健康碳水化合物來源有何說明？",
    "blind-32": "目前全球哪些疾病負擔促使 WHO 制定總脂肪與不健康體重增加的指引？",
}

INK = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
WHITE = "FFFFFF"
MUTED = "666666"
GRID = "B7C9DC"


def set_font(run, name, size, bold=False, color="000000", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph(paragraph, before=0, after=5, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_text(paragraph, text, font, size=10.5, bold=False, color="000000", italic=False):
    run = paragraph.add_run(text)
    set_font(run, font, size, bold, color, italic)
    return run


def add_page_field(paragraph, font):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "Page " if font == "Arial" else "第 ", font, 9, color=MUTED)
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])
    set_font(run, font, 9, color=MUTED)


def configure_document(doc, lang):
    font = "Arial" if lang == "en" else "Microsoft JhengHei"
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font
    normal._element.rPr.rFonts.set(qn("w:ascii"), font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, INK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:ascii"), font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(p, after=0)
    label = (
        "R2.22 Human Evidence Review | Blinded reviewer copy"
        if lang == "en"
        else "R2.22 真人證據盲審｜評審填寫版本"
    )
    add_text(p, label, font, 8.5, bold=True, color=MUTED)
    footer = section.footer
    add_page_field(footer.paragraphs[0], font)
    return font


def add_title_block(doc, font, lang):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=18, after=6)
    title = (
        "Blinded Human Evidence Review Workbook"
        if lang == "en"
        else "真人證據盲審評估表"
    )
    add_text(p, title, font, 24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=18)
    subtitle = (
        "32 nutrition and public-health questions | R2.22"
        if lang == "en"
        else "32 題營養與公共衛生問題｜R2.22"
    )
    add_text(p, subtitle, font, 12, color=MUTED)

    table = doc.add_table(rows=6, cols=2)
    set_table_geometry(table, [2100, 7260])
    fields = (
        [
            ("Reviewer code", "____________________________"),
            ("Professional qualification", "____________________________"),
            ("Current role / practice setting", "____________________________"),
            ("Years of relevant experience", "____________________________"),
            ("Review date", "____________________________"),
            ("Start and finish time", "____________________________"),
        ]
        if lang == "en"
        else [
            ("評審代碼", "____________________________"),
            ("專業資格", "____________________________"),
            ("目前職務／執業場域", "____________________________"),
            ("相關工作年資", "____________________________"),
            ("評估日期", "____________________________"),
            ("開始與完成時間", "____________________________"),
        ]
    )
    for row, (label, value) in zip(table.rows, fields):
        prevent_row_split(row)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, label, font, 9.5, bold=True, color=INK)
        p = row.cells[1].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, value, font, 10)


def add_instruction_sections(doc, font, lang):
    if lang == "en":
        sections = [
            (
                "1. What are we asking you to do?",
                [
                    "You will review 32 questions. Each question is followed by two source passages, labelled Candidate A and Candidate B.",
                    "Your task is not to grade the writing style and not to decide whether the World Health Organization is correct. Your task is to judge whether the supplied passages contain the evidence needed to answer the question.",
                    "Please judge only the text shown in this workbook. Do not use outside knowledge to fill a missing fact. If a passage is clipped, unclear, or affected by OCR errors, mention this in the rationale.",
                ],
            ),
            (
                "2. Simple review procedure",
                [
                    "Read the question first.",
                    "Read Candidate A and Candidate B in full.",
                    "Choose one answerability option and one evidence-contract option.",
                    "Assign one role to Candidate A and one role to Candidate B.",
                    "Choose a confidence score and briefly explain your reasoning.",
                    "Complete the items in order. Please do not compare your answers with another reviewer until all 32 items are finished.",
                ],
            ),
        ]
    else:
        sections = [
            (
                "一、我們希望您做什麼？",
                [
                    "本表共有 32 題。每題包含一個問題，以及標示為「候選段落 A」與「候選段落 B」的兩段來源文字。",
                    "本評估不是評分文字寫得好不好，也不是判斷世界衛生組織的立場是否正確。您的任務是判斷：僅依本表提供的段落，是否具有回答問題所需的證據。",
                    "請只依照本表顯示的文字判斷，不要用外部知識補上段落中缺少的資訊。若段落有截斷、語意不清或 OCR 亂碼，請在理由欄註明。",
                    "本表中的候選段落 A、B 為正式英文來源的繁體中文便利譯文。翻譯保留原文的數字、單位、條件與限制；若譯文或原始 OCR 內容仍有疑義，請勾選「文字品質疑慮」並在理由欄說明。",
                ],
            ),
            (
                "二、最簡單的填寫步驟",
                [
                    "先閱讀問題。",
                    "完整閱讀候選段落 A 與候選段落 B。",
                    "在「可回答程度」與「證據組合」各選一項。",
                    "分別判定候選段落 A 與候選段落 B 的角色。",
                    "選擇信心分數，並用一至三句話簡要說明理由。",
                    "請依題號順序完成。在 32 題全部完成前，請不要和其他評審對答案。",
                ],
            ),
        ]
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph(p, after=4)
            add_text(p, text, font, 10.3)

    doc.add_heading(
        "3. Detailed rating standards" if lang == "en" else "三、詳細評估標準",
        level=1,
    )
    criteria = (
        [
            ("Answerability", "Fully answerable", "The supplied passages contain all essential information requested by the question."),
            ("Answerability", "Partially answerable", "The passages answer part of the question, but at least one requested element is missing or unclear."),
            ("Answerability", "Not answerable", "The passages do not contain enough information to answer the central question."),
            ("Evidence contract", "Both required", "A and B each provide an essential requested element; removing either passage makes the answer incomplete."),
            ("Evidence contract", "A primary, B supporting", "A supplies the core answer. B adds useful interpretation, rationale, certainty, or context, but A still provides the main answer."),
            ("Evidence contract", "B primary, A supporting", "B supplies the core answer. A adds useful interpretation, rationale, certainty, or context."),
            ("Evidence contract", "A only / B only", "Only one passage is needed. The other may be related, repetitive, misleading, or irrelevant."),
            ("Evidence contract", "Neither", "Neither passage provides the evidence needed for the central question."),
            ("Candidate role", "Required", "Without this passage, an essential part of the requested answer is missing."),
            ("Candidate role", "Supporting", "The passage improves explanation or interpretation, but the core answer remains possible without it."),
            ("Candidate role", "Relevant but not required", "The passage is on-topic but redundant or unnecessary for this exact question."),
            ("Candidate role", "Unsafe or misleading", "Using the passage could produce a materially wrong answer because of wrong scope, outdated guidance, contradiction, or a misleading inference."),
            ("Candidate role", "Irrelevant", "The passage does not help answer this question."),
            ("Confidence", "1 to 5", "1 = very uncertain; 3 = moderately confident; 5 = very confident."),
        ]
        if lang == "en"
        else [
            ("可回答程度", "可完整回答", "所提供段落包含問題要求的所有必要資訊。"),
            ("可回答程度", "僅可部分回答", "段落能回答部分問題，但至少缺少一項問題要求的內容，或內容不清楚。"),
            ("可回答程度", "無法回答", "段落不足以回答問題的核心。"),
            ("證據組合", "A 與 B 都必要", "A 與 B 各提供一項不可缺少的內容；移除任一段都會使答案不完整。"),
            ("證據組合", "A 為主要、B 為補充", "A 提供核心答案；B 增加解釋、理由、確定性或背景，但主要答案仍來自 A。"),
            ("證據組合", "B 為主要、A 為補充", "B 提供核心答案；A 增加解釋、理由、確定性或背景。"),
            ("證據組合", "只需要 A／只需要 B", "只有其中一段是回答問題所必需；另一段可能只是相關、重複、誤導或無關。"),
            ("證據組合", "兩段都不需要", "A、B 都沒有提供回答核心問題所需的證據。"),
            ("候選段落角色", "必要", "若沒有此段，問題要求的重要部分就無法回答。"),
            ("候選段落角色", "補充", "此段能改善說明或解讀，但沒有它仍能回答核心問題。"),
            ("候選段落角色", "相關但非必要", "內容與主題有關，但對這個確切問題而言重複或不需要。"),
            ("候選段落角色", "可能造成不安全或誤導", "因適用範圍錯誤、過時、互相矛盾或易產生錯誤推論，使用此段可能造成實質錯誤。"),
            ("候選段落角色", "無關", "此段無助於回答本題。"),
            ("信心程度", "1 至 5 分", "1＝非常不確定；3＝中度確定；5＝非常確定。"),
        ]
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1850, 2260, 5250])
    headers = (
        ["Field", "Option", "How to decide"]
        if lang == "en"
        else ["評估欄位", "選項", "判斷方式"]
    )
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, text, font, 9.2, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    for field, option, rule in criteria:
        row = table.add_row()
        prevent_row_split(row)
        for index, text in enumerate((field, option, rule)):
            p = row.cells[index].paragraphs[0]
            set_paragraph(p, after=0, line=1.05)
            add_text(p, text, font, 9.0, bold=index == 1)

    doc.add_heading(
        "4. Worked examples" if lang == "en" else "四、判斷範例",
        level=1,
    )
    examples = (
        [
            ("Example A — both required", "Question asks for both benefit and safety. Passage A reports the benefit; Passage B reports the safety limitation. Choose “both required”; mark both candidates “required”."),
            ("Example B — one passage only", "Question asks for a numerical intake target. Passage A gives the target; Passage B only describes general background. Choose “A only”; mark A “required” and B “relevant but not required” or “irrelevant”."),
            ("Example C — primary plus supporting", "Question asks for a recommendation and how to interpret it. Passage A states the recommendation; Passage B explains low certainty. If A gives the core answer and B improves interpretation, choose “A primary, B supporting”."),
            ("Important distinction", "A passage is not “required” merely because it is useful or on-topic. Required means the answer would lose an essential requested element without it."),
        ]
        if lang == "en"
        else [
            ("範例 A－兩段都必要", "問題同時詢問效益與安全性。A 提供效益，B 提供安全限制。應選「A 與 B 都必要」，並將兩段都標為「必要」。"),
            ("範例 B－只需要一段", "問題詢問數值攝取目標。A 提供目標，B 只有一般背景。應選「只需要 A」；A 標為「必要」，B 標為「相關但非必要」或「無關」。"),
            ("範例 C－主要加補充", "問題詢問建議及其解讀方式。A 陳述建議，B 說明證據確定性偏低。若核心答案來自 A，而 B 改善解讀，應選「A 為主要、B 為補充」。"),
            ("重要區別", "一段文字不能只因為有幫助或與主題相關，就判定為「必要」。必要代表：拿掉這段後，答案會缺少問題明確要求的重要內容。"),
        ]
    )
    for label, text in examples:
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [2200, 7160])
        prevent_row_split(table.rows[0])
        set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
        p = table.rows[0].cells[0].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, label, font, 9.3, bold=True, color=INK)
        p = table.rows[0].cells[1].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, text, font, 9.3)

    doc.add_heading(
        "5. Reviewer declaration" if lang == "en" else "五、評審聲明",
        level=1,
    )
    declarations = (
        [
            "☐ I evaluated the items independently.",
            "☐ I did not see gold labels, strata, retrieval rankings, or experiment outcomes.",
            "☐ I judged only the supplied question and passages and did not use outside knowledge to repair missing evidence.",
            "☐ I will complete all 32 items before discussing answers with another person.",
        ]
        if lang == "en"
        else [
            "☐ 我是獨立完成本評估。",
            "☐ 我未看過 gold label、stratum、檢索排名或實驗結果。",
            "☐ 我只依本表提供的問題與段落判斷，未用外部知識補上缺漏證據。",
            "☐ 我會先完成全部 32 題，再與其他人討論答案。",
        ]
    )
    for text in declarations:
        p = doc.add_paragraph()
        set_paragraph(p, after=3)
        add_text(p, text, font, 10.3)
    p = doc.add_paragraph()
    set_paragraph(p, after=8)
    add_text(
        p,
        "Signature / initials: ____________________    Date: ____________________"
        if lang == "en"
        else "簽名／縮寫：____________________    日期：____________________",
        font,
        10.3,
    )
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=4, keep=True)
    add_text(
        p,
        "General notes or questions about the instructions (optional)"
        if lang == "en"
        else "對評估說明的整體備註或疑問（選填）",
        font,
        10.5,
        bold=True,
        color=BLUE,
    )
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    set_cell_shading(table.rows[0].cells[0], PALE)
    p = table.rows[0].cells[0].paragraphs[0]
    set_paragraph(p, after=0, line=1.35)
    add_text(
        p,
        "________________________________________________________________________\n"
        "________________________________________________________________________\n"
        "________________________________________________________________________\n"
        "________________________________________________________________________",
        font,
        10,
    )


def add_passage_table(doc, font, candidate, lang):
    table = doc.add_table(rows=2, cols=1)
    set_table_geometry(table, [9360])
    label = candidate["label"]
    heading = (
        f"Candidate {label} | {candidate['document_id']} | page {candidate['page_number']}"
        if lang == "en"
        else f"候選段落 {label}｜{candidate['document_id']}｜第 {candidate['page_number']} 頁｜中文便利譯文"
    )
    set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
    prevent_row_split(table.rows[0])
    p = table.rows[0].cells[0].paragraphs[0]
    set_paragraph(p, after=0)
    add_text(p, heading, font, 9.2, bold=True, color=INK)
    set_cell_shading(table.rows[1].cells[0], PALE)
    cell = table.rows[1].cells[0]
    cell.text = ""
    passage_text = candidate["text"] if lang == "en" else candidate["zh_text"]
    chunks = passage_text.splitlines() or [passage_text]
    passage_font = "Arial" if lang == "en" else font
    for index, chunk in enumerate(chunks):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_paragraph(p, after=3, line=1.03)
        add_text(p, chunk, passage_font, 8.4, color="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_assessment_table(doc, font, lang):
    rows = (
        [
            ("1. Answerability — choose one", "☐ Fully answerable    ☐ Partially answerable    ☐ Not answerable"),
            ("2. Evidence contract — choose one", "☐ Both required    ☐ A primary, B supporting    ☐ B primary, A supporting\n☐ A only    ☐ B only    ☐ Neither"),
            ("3. Candidate A role — choose one", "☐ Required    ☐ Supporting    ☐ Relevant but not required\n☐ Unsafe or misleading    ☐ Irrelevant"),
            ("4. Candidate B role — choose one", "☐ Required    ☐ Supporting    ☐ Relevant but not required\n☐ Unsafe or misleading    ☐ Irrelevant"),
            ("5. Confidence — choose one", "☐ 1 very uncertain    ☐ 2    ☐ 3 moderate    ☐ 4    ☐ 5 very confident"),
            ("6. Brief rationale", "____________________________________________________________________\n____________________________________________________________________\n____________________________________________________________________"),
            ("7. Text-quality concern (optional)", "☐ None    ☐ OCR problem    ☐ Clipped text    ☐ Other: __________________"),
        ]
        if lang == "en"
        else [
            ("1. 可回答程度－單選", "☐ 可完整回答    ☐ 僅可部分回答    ☐ 無法回答"),
            ("2. 證據組合－單選", "☐ A 與 B 都必要    ☐ A 為主要、B 為補充    ☐ B 為主要、A 為補充\n☐ 只需要 A    ☐ 只需要 B    ☐ 兩段都不需要"),
            ("3. 候選段落 A 的角色－單選", "☐ 必要    ☐ 補充    ☐ 相關但非必要\n☐ 可能造成不安全或誤導    ☐ 無關"),
            ("4. 候選段落 B 的角色－單選", "☐ 必要    ☐ 補充    ☐ 相關但非必要\n☐ 可能造成不安全或誤導    ☐ 無關"),
            ("5. 信心程度－單選", "☐ 1 非常不確定    ☐ 2    ☐ 3 中度確定    ☐ 4    ☐ 5 非常確定"),
            ("6. 簡短理由", "____________________________________________________________________\n____________________________________________________________________\n____________________________________________________________________"),
            ("7. 文字品質問題（選填）", "☐ 無    ☐ OCR 亂碼    ☐ 文字截斷    ☐ 其他：__________________"),
        ]
    )
    table = doc.add_table(rows=0, cols=2)
    for row_index, (label, options) in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p = row.cells[0].paragraphs[0]
        set_paragraph(p, after=0, line=1.05, keep=row_index < len(rows) - 1)
        add_text(p, label, font, 8.9, bold=True, color=INK)
        p = row.cells[1].paragraphs[0]
        set_paragraph(p, after=0, line=1.05, keep=row_index < len(rows) - 1)
        add_text(p, options, font, 8.9)
    set_table_geometry(table, [2600, 6760])


def add_item(doc, font, lang, item, index):
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.rows[0].cells[0], INK)
    p = table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(p, after=0)
    title = f"ITEM {index:02d} OF 32" if lang == "en" else f"第 {index:02d} 題／共 32 題"
    add_text(p, title, font, 12, bold=True, color=WHITE)

    question_label = "Question" if lang == "en" else "問題"
    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=2, keep=True)
    add_text(p, question_label, font, 10, bold=True, color=BLUE)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
    prevent_row_split(table.rows[0])
    p = table.rows[0].cells[0].paragraphs[0]
    set_paragraph(p, after=0, line=1.10)
    question = item["question"] if lang == "en" else ZH_QUESTIONS[item["blind_item_id"]]
    add_text(p, question, font, 11, bold=True, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    for candidate in item["candidates"]:
        add_passage_table(doc, font, candidate, lang)
    p = doc.add_paragraph()
    set_paragraph(p, before=3, after=3, keep=True)
    add_text(
        p,
        "Your assessment" if lang == "en" else "您的評估",
        font,
        10.5,
        bold=True,
        color=BLUE,
    )
    add_assessment_table(doc, font, lang)


def build_document(items, lang, output_path):
    doc = Document()
    font = configure_document(doc, lang)
    add_title_block(doc, font, lang)
    add_instruction_sections(doc, font, lang)
    for index, item in enumerate(items, 1):
        add_item(doc, font, lang, item, index)
    doc.core_properties.title = (
        "R2.22 Blinded Human Evidence Review Workbook"
        if lang == "en"
        else "R2.22 真人證據盲審評估表"
    )
    doc.core_properties.subject = "32-item blinded nutrition evidence review"
    doc.core_properties.author = "Healthy Diet AI Agent research project"
    doc.core_properties.keywords = "blinded review; nutrition; evidence; R2.22"
    doc.save(output_path)


def main():
    items = [json.loads(line) for line in PACKET.read_text(encoding="utf-8").splitlines() if line]
    if len(items) != 32:
        raise RuntimeError(f"Expected 32 blind items, got {len(items)}")
    if set(ZH_QUESTIONS) != {item["blind_item_id"] for item in items}:
        raise RuntimeError("Chinese question map does not match blind packet")
    translation_rows = [
        json.loads(line)
        for line in ZH_TRANSLATIONS.read_text(encoding="utf-8").splitlines()
        if line
    ]
    translations = {row["id"]: row for row in translation_rows}
    expected_translation_ids = {
        f"{item['blind_item_id']}-{candidate['label']}"
        for item in items
        for candidate in item["candidates"]
    }
    if len(translation_rows) != 64 or set(translations) != expected_translation_ids:
        raise RuntimeError("Chinese passage translation map must contain exactly 64 A/B records")
    for item in items:
        for candidate in item["candidates"]:
            translation_id = f"{item['blind_item_id']}-{candidate['label']}"
            translation = translations[translation_id]
            if translation["source_text"] != candidate["text"]:
                raise RuntimeError(f"Frozen source mismatch for {translation_id}")
            if not translation["zh_text"].strip():
                raise RuntimeError(f"Empty Chinese passage translation for {translation_id}")
            candidate["zh_text"] = translation["zh_text"]
    OUT.mkdir(parents=True, exist_ok=True)
    build_document(items, "en", OUT / "R2_22_HUMAN_REVIEW_EN.docx")
    build_document(items, "zh", OUT / "R2_22_HUMAN_REVIEW_ZH.docx")
    print(json.dumps({"status": "complete", "items": len(items), "output_dir": str(OUT)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
